import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from html2docx import html2docx
from .. import schemas
import requests
from io import BytesIO
import re
import markdown2
import html as html_lib

from xhtml2pdf import pisa

class ExportAgent:
    def _build_proposal_html(self, proposal: schemas.Proposal) -> str:
        """Builds the full HTML for a proposal, including CSS and content."""
        body_content = ""
        for section in sorted(proposal.sections, key=lambda s: s.order if s.order is not None else 999):
            raw = section.contentHtml or ""
            raw = html_lib.unescape(raw)
            if re.search(r"<\/?[a-zA-Z][\s\S]*?>", raw):
                html_from_markdown = raw
            else:
                html_from_markdown = markdown2.markdown(raw)

            mermaid_html = ""
            if section.mermaid_chart:
                mermaid_html = f'<div class="mermaid">{section.mermaid_chart}</div>'

            is_full_width = section.image_placement in ('full-width-top', 'full-width-bottom')
            image_html = ""
            if is_full_width and section.images:
                image_html = f'<div><img src="{section.images[0].url}" alt="{section.title}" style="width:100%; height:auto;" /></div>'

            text_wrapper = f"""
            <div style="padding: 2rem 3rem;">
                <h2 style="font-size:1.4rem; color:#2d3748; margin-bottom:0.75rem;">{section.title}</h2>
                <div class="content-wrapper" style="font-size:1rem; color:#333; line-height:1.6;">
                    {html_from_markdown}
                    {mermaid_html}
                </div>
            </div>
            """

            final_section_content = ""
            if is_full_width:
                if section.image_placement == 'full-width-top':
                    final_section_content = image_html + text_wrapper
                else:
                    final_section_content = text_wrapper + image_html
            else:
                final_section_content = text_wrapper

            layout_class = "two-column" if section.layout == 'two-column' else ""
            body_content += f'<div class="proposal-section {layout_class}" style="padding:0; overflow: hidden;">{final_section_content}</div>'

        html_content = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <title>Proposal for {proposal.clientName}</title>
            <style>
                body {{ font-family: sans-serif; }}
                .two-column .content-wrapper {{ column-count: 2; column-gap: 2rem; }}
                .proposal-container {{ max-width: 1000px; margin: 0 auto; background-color: white; }}
                {proposal.custom_css or ""}
            </style>
        </head>
        <body>
            <div class="proposal-container">
                <h1>Proposal for {proposal.clientName}</h1>
                {body_content}
            </div>
            <script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
            <script>mermaid.initialize({{startOnLoad:true}});</script>
        </body>
        </html>
        """
        return html_content

    def export_to_docx(self, proposal: schemas.Proposal) -> str:
        """Export proposal to a professional DOCX format with title page, headers, footers, and images."""
        document = Document()

        # --- Document Setup ---
        for section in document.sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

        # --- Header and Footer ---
        header = document.sections[0].header
        header_p = header.paragraphs[0]
        header_p.text = f"Proposal for {proposal.clientName}"
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        footer = document.sections[0].footer
        footer_p = footer.paragraphs[0]
        footer_p.text = f"{proposal.companyName} - Confidential"
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Title Page ---
        document.add_heading(f"Proposal for {proposal.clientName}", 0)
        document.add_paragraph()
        p_title = document.add_paragraph()
        p_title.add_run('Prepared for:').bold = True
        document.add_paragraph(proposal.clientName)
        document.add_paragraph()
        p_prepared_by = document.add_paragraph()
        p_prepared_by.add_run('Prepared by:').bold = True
        document.add_paragraph(proposal.companyName)
        document.add_paragraph(proposal.companyContact)
        document.add_paragraph(f"Date: {proposal.startDate.strftime('%B %d, %Y')}")
        document.add_page_break()

        # --- Main Content ---
        document.add_heading("Proposal Details", level=1)
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Item'
        hdr_cells[1].text = 'Details'
        details = {
            'Total Amount': f"${proposal.totalAmount:,.2f}",
            'Payment Type': proposal.paymentType,
            'Start Date': proposal.startDate.strftime('%Y-%m-%d'),
            'End Date': proposal.endDate.strftime('%Y-%m-%d'),
        }
        for key, value in details.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        document.add_paragraph()

        # --- Proposal Sections ---
        document.add_heading("Proposal Sections", level=1)
        for section in sorted(proposal.sections, key=lambda s: s.order):
            document.add_heading(section.title, level=2)

            image_url = section.images[0].url if section.images else None
            is_full_width = section.image_placement in ('full-width-top', 'full-width-bottom')

            def add_full_width_image(url):
                if not url: return
                try:
                    response = requests.get(url)
                    response.raise_for_status()
                    image_stream = BytesIO(response.content)
                    page_width = document.sections[-1].page_width - document.sections[-1].left_margin - document.sections[-1].right_margin
                    document.add_picture(image_stream, width=page_width)
                except Exception as e:
                    print(f"Error adding image from {url} to DOCX: {e}")

            if is_full_width and section.image_placement == 'full-width-top':
                add_full_width_image(image_url)

            try:
                html_to_render = section.contentHtml
                if not is_full_width and section.images:
                    for image in section.images:
                        html_to_render += f'<img src="{image.url}" />'
                if html_to_render:
                    html2docx(html_to_render, document)
            except Exception as e:
                print(f"Error parsing HTML for section '{section.title}': {e}")
                document.add_paragraph("Could not render section content.")

            if is_full_width and section.image_placement == 'full-width-bottom':
                add_full_width_image(image_url)

            document.add_page_break()

        # --- File Saving ---
        temp_dir = "backend/temp"
        os.makedirs(temp_dir, exist_ok=True)
        file_name = f"proposal_{proposal.id}.docx"
        file_path = os.path.join(temp_dir, file_name)
        document.save(file_path)
        return f"/exports/{file_name}"

    def export_to_pdf(self, proposal: schemas.Proposal) -> str:
        """Export proposal to a professional PDF format."""
        temp_dir = "backend/temp"
        os.makedirs(temp_dir, exist_ok=True)
        file_name = f"proposal_{proposal.id}.pdf"
        file_path = os.path.join(temp_dir, file_name)

        html_content = self._build_proposal_html(proposal)

        with open(file_path, "w+b") as pdf_file:
            pisa_status = pisa.CreatePDF(
                BytesIO(html_content.encode('UTF-8')),
                dest=pdf_file,
                encoding='UTF-8'
            )

        if pisa_status.err:
            print(f"Error creating PDF: {pisa_status.err}")
            return None

        return f"/exports/{file_name}"

export_agent = ExportAgent()